import asyncio
import json
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy.orm import sessionmaker
from docx import Document
import os

from app.core.config import settings
from app.db.database import Base
import app.models # Register all models
from app.models.user import User, UserProfile
from app.models.job import Job
from app.security.security import get_password_hash

async def seed_database():
    print("Connecting to database...")
    db_uri = settings.DATABASE_URI
    if db_uri.startswith("postgresql://"):
        db_uri = db_uri.replace("postgresql://", "postgresql+asyncpg://", 1)
    elif db_uri.startswith("postgres://"):
        db_uri = db_uri.replace("postgres://", "postgresql+asyncpg://", 1)

    engine = create_async_engine(db_uri, echo=False)
    
    # Create all tables (in a real app, use Alembic migrations)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    AsyncSessionLocal = sessionmaker(bind=engine, class_=AsyncSession, expire_on_commit=False)

    async with AsyncSessionLocal() as db:
        # Check if user exists
        from sqlalchemy.future import select
        result = await db.execute(select(User).filter(User.email == "test@example.com"))
        user = result.scalars().first()

        if not user:
            print("Creating test user: test@example.com / password123")
            user = User(
                email="test@example.com",
                hashed_password=get_password_hash("password123")
            )
            db.add(user)
            await db.commit()
            await db.refresh(user)

            print("Creating test profile...")
            profile = UserProfile(
                user_id=user.id,
                preferences_json={
                    "preferred_roles": ["Frontend Developer", "Full Stack Engineer"],
                    "preferred_locations": ["Remote", "New York"],
                    "remote_preference": "Remote",
                    "expected_salary_range": "$120k - $150k",
                    "notice_period_days": 14
                }
            )
            db.add(profile)
            await db.commit()

        # Create a test job
        result = await db.execute(select(Job).filter(Job.external_job_id == "TEST_JOB_1"))
        job = result.scalars().first()

        if not job:
            print("Creating test job...")
            job = Job(
                source="indeed",
                external_job_id="TEST_JOB_1",
                title="Senior Frontend Engineer (React/Next.js)",
                company="TechCorp Inc.",
                location="Remote",
                salary_range="$130k - $160k",
                url="https://example.com/jobs/1",
                description="We are looking for a Senior Frontend Engineer with deep expertise in React, Next.js, and TailwindCSS. You should have 5+ years of experience building scalable web applications. Experience with TypeScript and GraphQL is a huge plus. You will lead the frontend architecture for our new SaaS product."
            )
            db.add(job)
            await db.commit()

    print("Database seeded successfully!")

def create_sample_resume():
    """Generates a sample DOCX resume to be used for testing the parser."""
    os.makedirs("uploads/resumes", exist_ok=True)
    resume_path = "uploads/resumes/sample_resume.docx"
    
    if os.path.exists(resume_path):
        return

    doc = Document()
    doc.add_heading('John Doe', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Passionate Full Stack Developer with 4 years of experience building web applications using React, Python, and PostgreSQL. Proven ability to lead projects and mentor junior developers.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Languages: JavaScript, TypeScript, Python, HTML/CSS')
    doc.add_paragraph('Frameworks: React, Next.js, FastAPI, Node.js')
    doc.add_paragraph('Tools: Docker, Git, PostgreSQL, AWS')
    
    doc.add_heading('Experience', level=1)
    doc.add_heading('Software Engineer - StartupX (Jan 2021 - Present)', level=2)
    doc.add_paragraph('- Built the core MVP using Next.js and FastAPI, acquiring 10k users in the first month.')
    doc.add_paragraph('- Optimized database queries, reducing API latency by 40%.')
    
    doc.add_heading('Education', level=1)
    doc.add_heading('B.S. Computer Science - State University (2016 - 2020)', level=2)
    
    doc.save(resume_path)
    print(f"Sample resume generated at {resume_path}")

if __name__ == "__main__":
    create_sample_resume()
    asyncio.run(seed_database())
